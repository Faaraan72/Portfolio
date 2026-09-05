import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_bottom_border(paragraph, color_hex="00F0FF", size="12"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_hyperlink(paragraph, url, text, color_rgb=RGBColor(0, 150, 180), underline=True):
    """
    Adds a clickable XML hyperlink into a python-docx paragraph using the raw URL text.
    """
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Color
    c = OxmlElement('w:color')
    hex_color = f"{color_rgb[0]:02X}{color_rgb[1]:02X}{color_rgb[2]:02X}"
    c.set(qn('w:val'), hex_color)
    rPr.append(c)

    # Underline
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    # Font name & size
    f = OxmlElement('w:rFonts')
    f.set(qn('w:ascii'), 'Calibri')
    f.set(qn('w:hAnsi'), 'Calibri')
    rPr.append(f)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '19') # 9.5 pt
    rPr.append(sz)

    new_run.append(rPr)
    
    text_node = OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def build_resume():
    doc = docx.Document()
    
    # Standard margins (0.5 in top/bottom, 0.6 in left/right)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        
    # Color Palette
    COLOR_PRIMARY = RGBColor(10, 25, 47)      # Dark Navy
    COLOR_SECONDARY = RGBColor(0, 150, 180)   # Teal/Cyan Accent
    COLOR_MUTED = RGBColor(90, 100, 115)     # Grey for subtitles/dates
    COLOR_TEXT = RGBColor(30, 40, 50)        # Body text
    
    # ---------------------------------------------------------
    # HEADER SECTION
    # ---------------------------------------------------------
    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(2)
    p_name.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_name = p_name.add_run("MOHD FAARAAN ASKARI")
    r_name.font.name = "Calibri"
    r_name.font.size = Pt(22)
    r_name.font.bold = True
    r_name.font.color.rgb = COLOR_PRIMARY

    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(6)
    p_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("Unity & Multiplayer Systems Engineer | Netcode | Backend | XR")
    r_title.font.name = "Calibri"
    r_title.font.size = Pt(11)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_SECONDARY

    # Clean single-line contact info with raw URLs as clickable XML hyperlinks
    p_contact = doc.add_paragraph()
    p_contact.paragraph_format.space_before = Pt(0)
    p_contact.paragraph_format.space_after = Pt(12)
    p_contact.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    r_c1 = p_contact.add_run("+91 6307328745  |  mohdfaaraanaskari@gmail.com  |  ")
    r_c1.font.name = "Calibri"
    r_c1.font.size = Pt(9.5)
    r_c1.font.color.rgb = COLOR_MUTED

    add_hyperlink(p_contact, "https://linkedin.com/in/faaraan", "linkedin.com/in/faaraan", color_rgb=COLOR_SECONDARY)
    
    r_sep1 = p_contact.add_run("  |  ")
    r_sep1.font.name = "Calibri"
    r_sep1.font.size = Pt(9.5)
    r_sep1.font.color.rgb = COLOR_MUTED

    add_hyperlink(p_contact, "https://github.com/Faaraan72", "github.com/Faaraan72", color_rgb=COLOR_SECONDARY)

    r_sep2 = p_contact.add_run("  |  ")
    r_sep2.font.name = "Calibri"
    r_sep2.font.size = Pt(9.5)
    r_sep2.font.color.rgb = COLOR_MUTED

    add_hyperlink(p_contact, "https://faaraan72.github.io/Portfolio", "faaraan72.github.io/Portfolio", color_rgb=COLOR_SECONDARY)

    def add_section_heading(title_text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(title_text.upper())
        r.font.name = "Calibri"
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = COLOR_PRIMARY
        add_bottom_border(p, color_hex="008080", size="8")
        return p

    # ---------------------------------------------------------
    # PROFESSIONAL SUMMARY
    # ---------------------------------------------------------
    add_section_heading("Professional Summary")
    p_sum = doc.add_paragraph()
    p_sum.paragraph_format.space_before = Pt(0)
    p_sum.paragraph_format.space_after = Pt(8)
    p_sum.paragraph_format.line_spacing = 1.15
    r_sum = p_sum.add_run(
        "Unity & Multiplayer Systems Engineer with proven experience architecting real-time netcode (Photon Fusion, Nakama, NGO), containerized backend microservices (Docker, Edgegap, Python, MongoDB), and XR applications (Meta Quest VR, AR Foundation). Track record of shipping published mobile titles with 10k+ downloads and orchestrating multiplayer infrastructure supporting 100+ CCU at <120ms latency."
    )
    r_sum.font.name = "Calibri"
    r_sum.font.size = Pt(10)
    r_sum.font.color.rgb = COLOR_TEXT

    # ---------------------------------------------------------
    # TECHNICAL COMPETENCIES
    # ---------------------------------------------------------
    add_section_heading("Technical Competencies")
    
    skills = [
        ("Multiplayer & Netcode: ", "Photon Fusion (Host/Client & Shared), Nakama Server, Unity NGO, Client-Side Prediction, Server Reconciliation, State Synchronization, RPC Architecture."),
        ("Backend & DevOps: ", "Docker (Game Server Containerization), Edgegap (Global Orchestration), Python (FastAPI/Flask), MongoDB, Firebase (Auth, Firestore, Realtime DB, Analytics), WebSockets."),
        ("Game Engine Core & C#: ", "Unity 3D/2D, C#, Object Pooling, Event-Driven Architecture, Memory Profiling, Addressables, AssetBundles, Custom Physics Simulation."),
        ("Platforms & XR: ", "Meta Quest VR, Meta Movement SDK, AR Foundation, Vuforia, Mobile Deployment (iOS Xcode & Android Gradle), Windows PC."),
        ("Monetization & Analytics: ", "IronSource LevelPlay, Unity Ads, Unity IAP (SKU Mapping), Apple ATT Compliance, Google Play Billing.")
    ]
    
    for category, text in skills:
        p_sk = doc.add_paragraph()
        p_sk.paragraph_format.space_before = Pt(0)
        p_sk.paragraph_format.space_after = Pt(3)
        p_sk.paragraph_format.line_spacing = 1.1
        r_cat = p_sk.add_run(category)
        r_cat.font.name = "Calibri"
        r_cat.font.size = Pt(9.5)
        r_cat.font.bold = True
        r_cat.font.color.rgb = COLOR_PRIMARY
        
        r_txt = p_sk.add_run(text)
        r_txt.font.name = "Calibri"
        r_txt.font.size = Pt(9.5)
        r_txt.font.color.rgb = COLOR_TEXT

    # ---------------------------------------------------------
    # WORK EXPERIENCE (Commercial Production)
    # ---------------------------------------------------------
    add_section_heading("Work Experience")

    experiences = [
        {
            "company": "Shaurya Infosoft",
            "role": "Unity & Multiplayer Developer",
            "date": "June 2025 – Present",
            "bullets": [
                "Architecting a dedicated-server VR Multiplayer Tennis simulation for Meta Quest hardware using Docker containerization and Edgegap global server orchestration.",
                "Solo-engineered Tennis Freak 3D from initial prototype through production deployment for Android and iOS, managing gameplay physics, netcode sync, and cross-platform mobile deployment.",
                "Implemented Fusion multiplayer architecture for real-time state sync and custom matchmaking, supporting 50+ concurrent 1v1 matches (100+ CCU) at <120ms latency.",
                "Architected and integrated full-stack Python backend microservices and MongoDB database for server-side virtual currency processing and inventory persistence.",
                "Integrated Firebase (Google Play and Game Center Auth, Firestore, Analytics) across 1,000+ live test sessions.",
                "Developed Nakama-powered multiplayer titles (UNO, Connect 4), implementing custom matchmaking, synchronized turn logic, and cross-app data routing."
            ]
        },
        {
            "company": "Wizar Learning",
            "role": "Unity Developer",
            "date": "April 2024 – February 2025",
            "bullets": [
                "Engineered 10+ AR projects using Unity AR Foundation and Vuforia, enhancing user engagement through real-time interactive 3D features.",
                "Streamlined AR development pipelines and rendering passes using AR Foundation, boosting performance by 40% and ensuring frictionless deployment across Android and iOS.",
                "Collaborated with cross-functional technical teams to optimize asset streaming, texture compression, and Addressable asset memory footprints."
            ]
        },
        {
            "company": "Crimson Insights",
            "role": "Unity Developer Intern",
            "date": "Aug 2023 – Nov 2023",
            "bullets": [
                "Spearheaded end-to-end development of 'Zombie Attack', a 2D casual title published on Google Play (5K+ downloads, 4.0+ rating).",
                "Engineered core gameplay mechanics and character interactions through optimized C# scripts, boosting runtime execution speed by 40%.",
                "Orchestrated seamless integration of IronSource Ads, yielding a 20% increase in overall monetization revenue streams."
            ]
        },
        {
            "company": "Sagaci Studios",
            "role": "Unity 3D Developer Intern",
            "date": "May 2023 – June 2023",
            "bullets": [
                "Developed 'SoccerStriker', a mobile sports game reaching 10K+ downloads and 4.0+ rating on Google Play.",
                "Crafted real-world ball trajectory physics using C# and Unity Physics System, driving a 35% improvement in player session retention."
            ]
        },
        {
            "company": "HawkVisum Pvt Ltd",
            "role": "Unity 3D Developer Intern",
            "date": "Apr 2023 – June 2023",
            "bullets": [
                "Programmed C# simulation modules for a VR Aviation Cabin Crew Simulation project, writing 5,000+ lines of clean, modular C# code.",
                "Supervised a team of 3 3D modeling interns to ensure asset geometry, PBR materials, and draw calls adhered to VR performance budgets."
            ]
        }
    ]

    for exp in experiences:
        p_hdr = doc.add_paragraph()
        p_hdr.paragraph_format.space_before = Pt(6)
        p_hdr.paragraph_format.space_after = Pt(1)
        
        r_comp = p_hdr.add_run(exp["company"] + "  |  ")
        r_comp.font.name = "Calibri"
        r_comp.font.size = Pt(10.5)
        r_comp.font.bold = True
        r_comp.font.color.rgb = COLOR_PRIMARY
        
        r_role = p_hdr.add_run(exp["role"])
        r_role.font.name = "Calibri"
        r_role.font.size = Pt(10)
        r_role.font.italic = True
        r_role.font.color.rgb = COLOR_SECONDARY
        
        r_space = p_hdr.add_run("\t" + exp["date"])
        r_space.font.name = "Calibri"
        r_space.font.size = Pt(9.5)
        r_space.font.bold = True
        r_space.font.color.rgb = COLOR_MUTED

        for b in exp["bullets"]:
            p_b = doc.add_paragraph(style='List Bullet')
            p_b.paragraph_format.space_before = Pt(0)
            p_b.paragraph_format.space_after = Pt(2)
            p_b.paragraph_format.line_spacing = 1.15
            
            bullet_text = b.strip()
            if not bullet_text.endswith('.'):
                bullet_text += '.'

            r_b = p_b.add_run(bullet_text)
            r_b.font.name = "Calibri"
            r_b.font.size = Pt(9.5)
            r_b.font.color.rgb = COLOR_TEXT

    # ---------------------------------------------------------
    # FEATURED ENGINEERING PROJECTS (Enriched Technical Highlights)
    # ---------------------------------------------------------
    add_section_heading("Featured Projects")

    projects = [
        {
            "title": "CapsBall Multiplayer Mobile",
            "tech": "Unity, C#, Photon Fusion, Object Pooling, Event-Driven Architecture",
            "bullets": [
                "Engineered real-time physics-based multiplayer football using event-driven C# architecture (delegates/actions) with server-authoritative spawn and despawn state cleanup.",
                "Implemented high-performance Object Pooling pipelines for netcode network objects and dynamic particle FX to minimize garbage collection allocations and eliminate frame spikes.",
                "Integrated client-side state prediction and latency mitigation via entity interpolation for room timers and synchronized scores."
            ]
        },
        {
            "title": "Operation Air Siege - PC Flight Simulator",
            "tech": "Unity 3D, C#, Aerodynamics Physics, ScriptableObjects, Windows PC",
            "bullets": [
                "Developed a tactical flight combat simulator featuring aerodynamic lift and torque physics models, ScriptableObject weapon data pipelines, and missile guidance algorithms.",
                "Built event-driven observer architecture for audio/UI events, dynamic camera tracking loops, and custom vehicle control inputs optimized for 60+ FPS PC flight simulation."
            ]
        },
        {
            "title": "DeathMatch PC Tactical Shooter",
            "tech": "Unity 3D, C#, NavMesh AI, Object Pooling, Custom Shaders, Windows PC",
            "bullets": [
                "Built a third-person tactical shooter utilizing squad AI behaviors (NavMesh), weapon recoil/spread algorithms, and Object Pooling for projectile and effect instantiation.",
                "Designed high-fidelity combat visual effects, custom PBR shader graphs, and decoupled event-driven game state management."
            ]
        }
    ]

    for prj in projects:
        p_p = doc.add_paragraph()
        p_p.paragraph_format.space_before = Pt(6)
        p_p.paragraph_format.space_after = Pt(1)
        
        r_t = p_p.add_run(prj["title"] + "  ")
        r_t.font.name = "Calibri"
        r_t.font.size = Pt(10.5)
        r_t.font.bold = True
        r_t.font.color.rgb = COLOR_PRIMARY
        
        r_tch = p_p.add_run("(" + prj["tech"] + ")")
        r_tch.font.name = "Calibri"
        r_tch.font.size = Pt(9)
        r_tch.font.italic = True
        r_tch.font.color.rgb = COLOR_MUTED

        for b in prj["bullets"]:
            p_b = doc.add_paragraph(style='List Bullet')
            p_b.paragraph_format.space_before = Pt(0)
            p_b.paragraph_format.space_after = Pt(2)
            p_b.paragraph_format.line_spacing = 1.15
            
            bullet_text = b.strip()
            if not bullet_text.endswith('.'):
                bullet_text += '.'

            r_b = p_b.add_run(bullet_text)
            r_b.font.name = "Calibri"
            r_b.font.size = Pt(9.5)
            r_b.font.color.rgb = COLOR_TEXT

    # ---------------------------------------------------------
    # EDUCATION
    # ---------------------------------------------------------
    add_section_heading("Education")

    edus = [
        ("College of Engineering Roorkee", "Bachelor of Technology in Computer Science", "August 2020 – June 2024", "Roorkee, Uttarakhand"),
        ("City Montessori School", "Intermediate in Science (PCM + Computer Science)", "June 2017 – March 2019", "Lucknow, Uttar Pradesh")
    ]

    for inst, deg, dt, loc in edus:
        p_e = doc.add_paragraph()
        p_e.paragraph_format.space_before = Pt(4)
        p_e.paragraph_format.space_after = Pt(2)
        r_i = p_e.add_run(inst + "  |  ")
        r_i.font.name = "Calibri"
        r_i.font.size = Pt(10)
        r_i.font.bold = True
        r_i.font.color.rgb = COLOR_PRIMARY
        
        r_d = p_e.add_run(deg)
        r_d.font.name = "Calibri"
        r_d.font.size = Pt(9.5)
        r_d.font.italic = True
        r_d.font.color.rgb = COLOR_TEXT
        
        r_dt = p_e.add_run("\t" + dt)
        r_dt.font.name = "Calibri"
        r_dt.font.size = Pt(9)
        r_dt.font.color.rgb = COLOR_MUTED

    output_path = "c:\\Users\\faaraan\\OneDrive\\Desktop\\github\\Portfolio\\Mohd_Faaraan_Askari_CV.docx"
    try:
        doc.save(output_path)
        print(f"CV successfully generated at: {output_path}")
    except PermissionError:
        fallback_path = "c:\\Users\\faaraan\\OneDrive\\Desktop\\github\\Portfolio\\Mohd_Faaraan_Askari_CV_Updated.docx"
        doc.save(fallback_path)
        print(f"CV successfully saved to fallback path: {fallback_path}")

if __name__ == "__main__":
    build_resume()
